import os
import uuid
from typing import List, Dict, Tuple, Optional
import hashlib
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import pytesseract

# LangChain components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_community.vectorstores import FAISS

from config import Config
from database import db

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separators=["\n\n", "\n", " ", ""]
        )
        # Use vector store manager for embeddings (supports both Google and Cohere)
        from vector_store_manager import vector_store_manager
        self.vector_store_manager = vector_store_manager
        self.vector_stores = {}  # Cache for vector stores

    def process_uploaded_file(self, uploaded_file) -> Tuple[int, str]:
        """Process an uploaded file and return document_id and status message"""
        try:
            # Generate unique filename
            file_extension = uploaded_file.name.split('.')[-1].lower()
            unique_filename = f"{uuid.uuid4()}.{file_extension}"

            # Save file temporarily
            temp_path = f"temp_{unique_filename}"
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getvalue())

            # Process document based on file type
            documents = self.extract_text_from_file(temp_path, file_extension)

            if not documents:
                os.remove(temp_path)
                return None, "Could not extract text from the document"

            # Add to database
            document_id = db.add_document(
                filename=unique_filename,
                original_filename=uploaded_file.name,
                file_type=file_extension,
                file_size=uploaded_file.size,
                description=f"Uploaded {file_extension.upper()} file"
            )

            # Process text and create vector store
            chunks = self.text_splitter.split_documents(documents)

            # Add chunks to database for tracking
            chunk_texts = [chunk.page_content for chunk in chunks]
            chunk_metadata = [chunk.metadata for chunk in chunks]
            db.add_document_chunks(document_id, chunk_texts, chunk_metadata)

            # Create vector store using vector_store_manager
            from vector_store_manager import vector_store_manager
            vector_store_path = vector_store_manager.create_vector_store(document_id, chunks)

            # Update document status
            db.update_document_processed(document_id, vector_store_path)

            # Clean up temp file
            os.remove(temp_path)

            return document_id, f"Successfully processed {uploaded_file.name} with {len(chunks)} chunks"

        except Exception as e:
            # Clean up temp file if it exists
            if 'temp_path' in locals() and os.path.exists(temp_path):
                os.remove(temp_path)
            return None, f"Error processing file: {str(e)}"

    def extract_text_from_file(self, file_path: str, file_type: str) -> List[LangchainDocument]:
        """Extract text from various file types"""
        try:
            if file_type == 'pdf':
                return self.extract_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                return self.extract_from_docx(file_path)
            elif file_type in ['txt', 'md']:
                return self.extract_from_text(file_path)
            elif file_type in ['csv', 'xlsx', 'xls']:
                return self.extract_from_spreadsheet(file_path, file_type)
            elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
                return self.extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text from {file_type}: {str(e)}")

    def extract_from_pdf(self, file_path: str) -> List[LangchainDocument]:
        """Extract text from PDF file"""
        documents = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    metadata = {
                        'source': file_path,
                        'page': page_num + 1,
                        'file_type': 'pdf'
                    }
                    documents.append(LangchainDocument(page_content=text, metadata=metadata))
        return documents

    def extract_from_docx(self, file_path: str) -> List[LangchainDocument]:
        """Extract text from DOCX file"""
        documents = []
        doc = DocxDocument(file_path)
        full_text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        if full_text:
            text = '\n'.join(full_text)
            metadata = {
                'source': file_path,
                'file_type': 'docx'
            }
            documents.append(LangchainDocument(page_content=text, metadata=metadata))

        return documents

    def extract_from_text(self, file_path: str) -> List[LangchainDocument]:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

        metadata = {
            'source': file_path,
            'file_type': 'text'
        }

        return [LangchainDocument(page_content=text, metadata=metadata)]

    def extract_from_spreadsheet(self, file_path: str, file_type: str) -> List[LangchainDocument]:
        """Extract text from Excel or CSV file"""
        if file_type == 'csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        # Convert dataframe to text
        text = df.to_string(index=False)

        metadata = {
            'source': file_path,
            'file_type': file_type,
            'columns': list(df.columns),
            'rows': len(df)
        }

        return [LangchainDocument(page_content=text, metadata=metadata)]

    def extract_from_image(self, file_path: str) -> List[LangchainDocument]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)

            if text.strip():
                metadata = {
                    'source': file_path,
                    'file_type': 'image',
                    'ocr_processed': True
                }
                return [LangchainDocument(page_content=text, metadata=metadata)]
            else:
                return []
        except Exception as e:
            raise Exception(f"OCR processing failed: {str(e)}. Make sure Tesseract OCR is installed.")

    def get_vector_store(self, document_id: int) -> Optional[FAISS]:
        """Get vector store for a document, loading from disk if necessary"""
        if document_id in self.vector_stores:
            return self.vector_stores[document_id]

        # Load from disk using vector store manager's embeddings
        vector_store_path = f"{Config.VECTOR_STORE_PATH}/{document_id}"
        if os.path.exists(vector_store_path):
            try:
                # Reinitialize embeddings to ensure current configuration
                self.vector_store_manager.reinitialize_if_needed()
                vector_store = FAISS.load_local(vector_store_path, self.vector_store_manager.embeddings, allow_dangerous_deserialization=True)
                self.vector_stores[document_id] = vector_store
                return vector_store
            except Exception as e:
                print(f"Error loading vector store for document {document_id}: {str(e)}")

        return None

    def search_similar_documents(self, document_id: int, query: str, k: int = 4) -> List[LangchainDocument]:
        """Search for similar documents in the vector store"""
        vector_store = self.get_vector_store(document_id)
        if not vector_store:
            return []

        try:
            results = vector_store.similarity_search(query, k=k)
            return results
        except Exception as e:
            print(f"Error searching documents: {str(e)}")
            return []

    def search_with_scores(self, document_id: int, query: str, k: int = 4) -> List[Tuple[LangchainDocument, float]]:
        """Search for similar documents with similarity scores"""
        vector_store = self.get_vector_store(document_id)
        if not vector_store:
            return []

        try:
            results = vector_store.similarity_search_with_score(query, k=k)
            return results
        except Exception as e:
            print(f"Error searching documents with scores: {str(e)}")
            return []

    def delete_document_vector_store(self, document_id: int):
        """Delete vector store for a document"""
        # Remove from cache
        if document_id in self.vector_stores:
            del self.vector_stores[document_id]

        # Delete from disk
        vector_store_path = f"{Config.VECTOR_STORE_PATH}/{document_id}"
        if os.path.exists(vector_store_path):
            import shutil
            shutil.rmtree(vector_store_path)

    def get_document_info(self, document_id: int) -> Dict:
        """Get comprehensive information about a document"""
        document = db.get_document(document_id)
        if not document:
            return {}

        chunks = db.get_document_chunks(document_id)

        info = {
            'document': document,
            'chunks_count': len(chunks),
            'vector_store_exists': os.path.exists(f"{Config.VECTOR_STORE_PATH}/{document_id}"),
            'total_characters': sum(len(chunk['content']) for chunk in chunks),
            'preview': chunks[0]['content'][:200] + "..." if chunks else "No content available"
        }

        return info

# Global document processor instance
document_processor = DocumentProcessor()